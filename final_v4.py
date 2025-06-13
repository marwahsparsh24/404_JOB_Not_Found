import os
import fitz
import chromadb
import pandas as pd
from pymongo import MongoClient
from sklearn.metrics.pairwise import cosine_similarity
import docx
import requests
import numpy as np
from collections import Counter
from openai import OpenAI
from crewai import Agent, Task, Crew
from datetime import datetime, timezone
import os
import re

os.environ["OPENAI_API_KEY"] = "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"
# 🔐 Load API key from environment variable
openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Define folder paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RESUME_DIR = os.path.join(BASE_DIR, 'data', 'resumes')
JD_DIR = os.path.join(BASE_DIR,'data','job_descriptions')
STATUS_DIR = os.path.join(BASE_DIR,'data','application_status')


def extract_formatted_text(file_path):
    """
    Extract formatted text from PDF or DOCX files.
    For PDFs: preserve font styles, sizes, layout using PyMuPDF.
    For DOCX: extract paragraph-wise text without formatting details.
    """
    try:
        if file_path.lower().endswith(".pdf"):
            doc = fitz.open(file_path)
            formatted_text = ""

            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                blocks = page.get_text("dict")["blocks"]

                for block in blocks:
                    if block["type"] == 0:  # Only process text blocks
                        for line in block["lines"]:
                            for span in line["spans"]:
                                text = span["text"]
                                font_size = span["size"]
                                font_color = span["color"]
                                font_name = span["font"]
                                is_bold = "bold" in font_name.lower()
                                is_italic = "italic" in font_name.lower()

                                formatted_text += (
                                    f"Text: {text}, Font: {font_name}, Size: {font_size}, "
                                    f"Color: {font_color}, Bold: {is_bold}, Italic: {is_italic}\n"
                                )
            return formatted_text

        elif file_path.lower().endswith(".docx"):
            doc = docx.Document(file_path)
            formatted_text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    formatted_text += f"Text: {para.text}, Font: N/A, Size: N/A, Color: N/A, Bold: N/A, Italic: N/A\n"
            return formatted_text

        else:
            return "Unsupported file format."

    except Exception as e:
        return f"Error extracting formatted text from {file_path}: {e}"


def extract_plain_text(pdf_path):
    """Extract plain text (ATS-friendly) using PyMuPDF."""
    try:
        doc = fitz.open(pdf_path)
        plain_text = ""

        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            plain_text += page.get_text("text")

        return plain_text
    except Exception as e:
        print(f"Error extracting plain text from {pdf_path}: {e}")
        return ""


def read_text_file(file_path):
    """Read text from JD or application status files."""
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read().strip()
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return ""


def process_resumes():
    """Processes resumes, JDs, and application statuses into a Pandas DataFrame."""
    data = []

    # List all files in the directories
    resume_files = sorted([file for file in os.listdir(RESUME_DIR) if file.lower().endswith('.pdf')])
    jd_files = sorted([file for file in os.listdir(JD_DIR) if file.lower().endswith('.txt')])
    status_files = sorted([file for file in os.listdir(STATUS_DIR) if file.lower().endswith('.txt')])

    # Ensure that the lengths of resume_files, jd_files, and status_files are consistent
    num_files = min(len(resume_files), len(jd_files), len(status_files))

    for i in range(num_files):
        resume_file = resume_files[i]
        jd_file = jd_files[i]
        status_file = status_files[i]

        resume_path = os.path.join(RESUME_DIR, resume_file)
        jd_path = os.path.join(JD_DIR, jd_file)
        status_path = os.path.join(STATUS_DIR, status_file)

        formatted_text = extract_formatted_text(resume_path)
        plain_text = extract_plain_text(resume_path)
        job_description = read_text_file(jd_path)
        application_status = read_text_file(status_path)

        # Store everything in a structured list
        data.append({
            "Resume File": resume_file,
            "Plain Text Resume": plain_text,
            "Formatted Resume": formatted_text,
            "Job Description": job_description,
            "Application Status": application_status
        })

    # Convert to DataFrame
    df = pd.DataFrame(data)

    return df

# Run the script and get DataFrame
df = process_resumes()
# Display DataFrame


mongo_client = MongoClient("mongodb://localhost:27017/")
db = mongo_client["resume_db"]
formatted_resume_collection = db["formatted_resumes"]

client = chromadb.PersistentClient(path="./chroma_db")
collection = client.get_or_create_collection("plain_resumes")


def store_to_chromadb_and_mongodb(df):
    for index, row in df.iterrows():
        resume_file = row["Resume File"]
        plain_text = row["Plain Text Resume"]
        formatted_text = row["Formatted Resume"]
        job_description = row["Job Description"]
        application_status = row["Application Status"]

        # MongoDB
        existing_resume = formatted_resume_collection.find_one({"resume_file": resume_file})
        if existing_resume:
            formatted_resume_collection.update_one(
                {"resume_file": resume_file},
                {"$set": {
                    "formatted_resume": formatted_text,
                    "job_description": job_description,
                    "application_status": application_status
                }}
            )
            print(f"Updated existing entry in MongoDB for: {resume_file}")
        else:
            formatted_resume_collection.insert_one({
                "resume_file": resume_file,
                "formatted_resume": formatted_text,
                "job_description": job_description,
                "application_status": application_status
            })
            print(f"Inserted new entry in MongoDB for: {resume_file}")

        # === Store Resume Embedding in ChromaDB ===
        resume_doc_id = f"resume_{resume_file}"
        if collection.get(ids=[resume_doc_id])['documents']:
            collection.delete(ids=[resume_doc_id])
        collection.add(
            documents=[plain_text],
            metadatas=[{
                "type": "resume",
                "resume_file": resume_file,
                "job_description": job_description,
                "application_status": application_status
            }],
            ids=[resume_doc_id]
        )
        print(f"Stored in ChromaDB: {resume_doc_id}")

        # === Store JD Embedding in ChromaDB ===
        jd_doc_id = f"jd_{resume_file}"
        if collection.get(ids=[jd_doc_id])['documents']:
            collection.delete(ids=[jd_doc_id])
        collection.add(
            documents=[job_description],
            metadatas=[{
                "type": "job_description",
                "resume_file": resume_file,
                "application_status": application_status
            }],
            ids=[jd_doc_id]
        )
        print(f"Stored in ChromaDB: {jd_doc_id}")

store_to_chromadb_and_mongodb(df)

# MongoDB connection string (adjust it as needed)
MONGO_URI = "mongodb://localhost:27017"
DATABASE_NAME = "resume_db"
FORMATTED_COLLECTION = "formatted_resumes"

# Connect to MongoDB
client = MongoClient(MONGO_URI)
db = client[DATABASE_NAME]
formatted_resume_collection = db[FORMATTED_COLLECTION]


def retrieve_formatted_resumes_as_df():
    """Retrieve unique formatted resume data from MongoDB and return it as a DataFrame"""

    # Retrieve all documents from the collection
    formatted_resumes = formatted_resume_collection.find()

    # Convert the results into a list of dictionaries while ensuring unique entries
    resume_list = []
    seen_files = set()  # To track unique resume files

    for resume in formatted_resumes:
        resume_file = resume['resume_file']

        # Only append if the resume_file hasn't been encountered before
        if resume_file not in seen_files:
            seen_files.add(resume_file)
            resume_list.append({
                "Resume File": resume['resume_file'],
                "Formatted Resume": resume['formatted_resume'],
                "Job Description": resume['job_description'],
                "Application Status": resume['application_status']
            })

    # Convert the list of dictionaries to a DataFrame
    df = pd.DataFrame(resume_list)
    return df


# Example usage: retrieve unique formatted resumes as a DataFrame
formatted_resume_df = retrieve_formatted_resumes_as_df()
print(formatted_resume_df)
# Initialize ChromaDB client and collection
client = chromadb.PersistentClient(path="./chroma_db")  # Use the correct path if needed
collection = client.get_or_create_collection(name="plain_resumes")


def retrieve_all_resumes():
    """Retrieve all resumes stored in ChromaDB and return as a DataFrame"""
    results = collection.get()  # Retrieves all stored documents

    # Convert results to a structured DataFrame
    resume_list = []
    for i in range(len(results['documents'])):
        resume_list.append({
            "Resume File": results['ids'][i],
            "Plain Text Resume": results['documents'][i],
            "Metadata": results['metadatas'][i]  # Contains metadata like resume_file, job_description, etc.
        })

    df = pd.DataFrame(resume_list)
    return df


# Retrieve all resumes
plain_resumes_df = retrieve_all_resumes()


INPUT_DIR = os.path.join(BASE_DIR,'data', 'inputdata')
os.makedirs(INPUT_DIR, exist_ok=True)

# Initialize empty DataFrames if not already defined
try:
    formatted_resume_df
except NameError:
    formatted_resume_df = pd.DataFrame(columns=["Resume File", "Formatted Resume", "Job Description", "Application Status"])

try:
    plain_resumes_df
except NameError:
    plain_resumes_df = pd.DataFrame(columns=["Resume File", "Plain Text Resume", "Metadata"])

# Save uploaded file to disk (used by backend functions)

def save_uploaded_file(file_name, file_bytes):
    """Save uploaded resume to disk in the input folder."""
    file_path = os.path.join(INPUT_DIR, file_name)
    with open(file_path, "wb") as f:
        f.write(file_bytes)
    return file_name  # Just returning the filename

# === Core Function ===
def upload_rejected_resume_local_to_df(resume_filename, jd_text, application_status):
    if application_status.lower() != "rejected":
        return {"error": "Application status must be 'rejected'"}

    if not jd_text.strip():
        return {"error": "Job description text cannot be empty."}

    resume_path = os.path.join(INPUT_DIR, resume_filename)

    formatted_text = extract_formatted_text(resume_path)
    plain_text = extract_plain_text(resume_path)

    formatted_row = {
        "Resume File": resume_filename,
        "Formatted Resume": formatted_text,
        "Job Description": jd_text,
        "Application Status": application_status
    }

    plain_row = {
        "Resume File": f"resume_{resume_filename}",
        "Plain Text Resume": plain_text,
        "Metadata": {
            "application_status": application_status,
            "job_description": jd_text,
            "resume_file": resume_filename,
        }
    }

    jd_row = {
        "Resume File": f"jd_{resume_filename}",
        "Plain Text Resume": jd_text,
        "Metadata": {
            "application_status": application_status,
            "resume_file": resume_filename,
            "type": "job_description"
        }
    }

    global formatted_resume_df, plain_resumes_df
    formatted_resume_df = formatted_resume_df[formatted_resume_df["Resume File"] != resume_filename]
    plain_resumes_df = plain_resumes_df[
        ~plain_resumes_df["Resume File"].isin([f"resume_{resume_filename}", f"jd_{resume_filename}"])]

    formatted_resume_df = pd.concat([formatted_resume_df, pd.DataFrame([formatted_row])], ignore_index=True)
    plain_resumes_df = pd.concat([plain_resumes_df, pd.DataFrame([plain_row, jd_row])], ignore_index=True)

    # === Store in MongoDB ===
    existing_resume = formatted_resume_collection.find_one({"resume_file": resume_filename})
    if existing_resume:
        formatted_resume_collection.update_one(
            {"resume_file": resume_filename},
            {"$set": {
                "formatted_resume": formatted_text,
                "job_description": jd_text,
                "application_status": application_status
            }}
        )
    else:
        formatted_resume_collection.insert_one({
            "resume_file": resume_filename,
            "formatted_resume": formatted_text,
            "job_description": jd_text,
            "application_status": application_status
        })

    # === Store in ChromaDB ===
    resume_doc_id = f"resume_{resume_filename}"
    jd_doc_id = f"jd_{resume_filename}"

    if collection.get(ids=[resume_doc_id])["documents"]:
        collection.delete(ids=[resume_doc_id])
    if collection.get(ids=[jd_doc_id])["documents"]:
        collection.delete(ids=[jd_doc_id])

    collection.add(
        documents=[plain_text],
        metadatas=[{
            "type": "resume",
            "resume_file": resume_filename,
            "job_description": jd_text,
            "application_status": application_status
        }],
        ids=[resume_doc_id]
    )

    collection.add(
        documents=[jd_text],
        metadatas=[{
            "type": "job_description",
            "resume_file": resume_filename,
            "application_status": application_status
        }],
        ids=[jd_doc_id]
    )

    return {
        "message": "Files processed and stored in local DataFrames, MongoDB, and ChromaDB (replaced if existed)",
        "resume_file": resume_filename
    }


# === 1. FORMATTING COMPLIANCE CHECK ===
def run_formatting_compliance_last():
    global formatted_resume_df

    if formatted_resume_df.empty:
        return "⚠️ DataFrame is empty."

    last_row = formatted_resume_df.iloc[-1]
    formatted_text = last_row["Formatted Resume"]
    resume_file = last_row["Resume File"]

    system_prompt = (
        "You are an expert in resume formatting for Applicant Tracking Systems (ATS). "
        "Your task is to analyze formatting features extracted from resumes and identify ATS-related issues."
    )

    user_prompt = ats_prompt = f"""
    You are an expert ATS (Applicant Tracking System) evaluator and resume coach.
    
    Your task is to evaluate the following resume in detail for ATS compatibility and effectiveness, based on best practices from tools like Jobscan, Resumeworded, and major employers' ATS systems.
    
    === RESUME TEXT ===
    {formatted_text}
    
    Please evaluate the resume using these criteria and return clear, structured feedback:
    
    ---
    
    1. **Font & Formatting Compliance**
    - Is the font ATS-safe (e.g., Calibri, Arial, Times New Roman)?
    - Are the font sizes consistent (10–12pt for body, 12–14pt headers)?
    - Is the layout one-column, without tables, text boxes, columns, headers, footers, or images?
    
    2. **File Type & Layout Clarity**
    - Does the resume appear parseable by an ATS (no graphics, icons, or scanned content)?
    - Does it avoid fancy design elements and maintain clean spacing?
    
    3. **Keyword Optimization & Contextual Match**
    - Does the resume include **exact keywords** from the job description (skills, tools, job titles)?
    - Are both **hard** and **soft** skills present?
    - Is there good keyword distribution across the **Summary**, **Experience**, and **Skills**?
    
    4. **Section Title Clarity**
    - Are standard sections present with conventional labels like:
      - "Professional Experience", "Education", "Skills", "Certifications", "Summary"?
    - Flag if any non-standard or unclear section headings are used.
    
    5. **Bullet Usage & Structure**
    - Does the resume use clean bullet points instead of paragraphs?
    - Are bullets concise (1–2 lines), with consistent symbols like •?
    
    6. **Quantified Achievements**
    - Are specific results or numbers included (e.g., “Improved X by 20%”)?
    - If missing, suggest where they could be added.
    
    7. **Contact Information**
    - Is contact info (name, phone, email, LinkedIn) cleanly placed at the top?
    - Is it **outside headers/footers** for parseability?
    
    8. **Section Completeness**
    - Are required sections **(Experience, Education, Skills, Summary)** all included?
    - If any are missing, flag them clearly.
    
    9. **Skills Section Quality**
    - Are **hard skills** and tools included (e.g., Python, SQL, Tableau)?
    - If it lists soft skills here, suggest moving them elsewhere.
    
    10. **Consistency & Clean Structure**
    - Are dates consistently formatted (e.g., Jan 2020 – Mar 2022)?
    - Are all headings, job titles, and bullets properly aligned?
    
    11. **Resume Length Appropriateness**
    - One page for <5 years experience
    - Two pages for >5 years
    - If it's too long or short, suggest trimming or expanding.
    
    12. **Semantic Relevance & Variation**
    - Does the resume demonstrate understanding of the job by using natural language variation?
    - Example: Instead of repeating "Managed", does it use synonyms like “Led”, “Oversaw”, “Directed”?
    
    ---
    
    ✅ **Final Output Format**:
    - Short Summary
    - Bullet-based Section Feedback (one for each of the categories above)
    - Suggested Job Title (optional)
    - Clean formatting, no asterisks or special symbols
    - No placeholders like [Company Name]
    
    Be direct and specific. Write this as feedback for a serious jobseeker optimizing for ATS.
    
    """


    try:
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.2,
            max_tokens=500
        )
        report = response.choices[0].message.content.strip()
    except Exception as e:
        report = f"Error during analysis: {e}"

    print(f"✅ Formatting check completed for: {resume_file}")
    return report

#ATS Score
import re
from sklearn.metrics.pairwise import cosine_similarity

# Helper: Check for section presence
def evaluate_section_presence(resume_text):
    required_sections = ["experience", "education", "skills", "projects", "summary"]
    found = [sec for sec in required_sections if sec in resume_text.lower()]
    return round((len(found) / len(required_sections)) * 100, 2)

# Helper: Evaluate formatting quality
def evaluate_formatting_score(resume_text):
    penalties = 0
    if 'table' in resume_text.lower(): penalties += 1
    if resume_text.lower().count('\n\n') < 3: penalties += 1
    if len(re.findall(r'\|', resume_text)) > 5: penalties += 1  # Too many pipe symbols
    score = max(0, 100 - penalties * 20)
    return round(score, 2)

# Helper: Evaluate structure (length, bullet points, dates)
def evaluate_structure_score(resume_text):
    if len(resume_text) < 1000: return 40.0  # too short
    bullet_count = resume_text.count('•') + resume_text.count('- ')
    date_matches = re.findall(r'\b(19|20)\d{2}\b', resume_text)
    score = 50
    if bullet_count >= 5: score += 20
    if len(date_matches) >= 3: score += 20
    return min(score, 100)

# === MAIN: Enhanced ATS scoring ===
def run_ats_scoring_with_embedding():
    non_jd_df = plain_resumes_df[~plain_resumes_df["Resume File"].str.startswith("jd_")]
    if non_jd_df.empty:
        return "⚠️ No resumes to score."

    last_resume = non_jd_df.iloc[-1]
    resume_file = last_resume["Resume File"]
    resume_text = last_resume["Plain Text Resume"]
    jd_entry = plain_resumes_df[plain_resumes_df["Resume File"] == f"jd_{resume_file.split('_', 1)[-1]}"]

    if jd_entry.empty:
        return f"❌ No matching job description found for resume `{resume_file}`."

    jd_text = jd_entry.iloc[0]["Plain Text Resume"]

    # --- Score 1: Embedding similarity ---
    resume_embedding = embed_text(resume_text)
    jd_embedding = embed_text(jd_text)
    sim_score = cosine_similarity([resume_embedding], [jd_embedding])[0][0] * 100

    # --- Score 2: Keyword match ---
    resume_words = set(re.findall(r'\b\w+\b', resume_text.lower()))
    jd_keywords = set(extract_keywords(jd_text, limit=30))
    keyword_overlap = len(jd_keywords & resume_words)
    keyword_score = (keyword_overlap / len(jd_keywords)) * 100 if jd_keywords else 0

    # --- Score 3: Section presence ---
    section_score = evaluate_section_presence(resume_text)

    # --- Score 4: Formatting ---
    formatting_score = evaluate_formatting_score(resume_text)

    # --- Score 5: Style / structure ---
    structure_score = evaluate_structure_score(resume_text)

    # === Final weighted score ===
    final_score = round(
        0.30 * sim_score +
        0.25 * keyword_score +
        0.15 * section_score +
        0.15 * formatting_score +
        0.15 * structure_score,
        2
    )

    # === Reasoning Breakdown ===
    keyword_reasoning = f"There is a {'significant' if keyword_score > 50 else 'limited'} overlap between the resume and the job description keywords. Match score: {keyword_score:.2f}%."
    formatting_reasoning = "The resume is well-formatted with clear section headings and bullet points. It includes contact info and links to professional profiles."
    missing_section_reasoning = "The resume does not have a 'References' section, which may be required depending on employer preference. Not critical but worth noting."

    try:
        response = openai_client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an expert resume and job description evaluator."},
                {"role": "user", "content": f"""
Evaluate this resume against the job description.
Resume:
{resume_text[:4000]}

Job Description:
{jd_text[:4000]}

Highlight keyword overlap, formatting issues, missing sections, and give a short conclusion.
"""}
            ],
            temperature=0.3,
            max_tokens=500
        )
        conclusion_reasoning = response.choices[0].message.content.strip()
    except Exception as e:
        conclusion_reasoning = f"(LLM reasoning failed: {e})"
    matched_keywords = list(jd_keywords & resume_words)
    unmatched_keywords = list(jd_keywords - resume_words)
    return {
        "score": final_score,
        "resume_file": resume_file,
        "keywords": {  # Add this new section
            "matched": matched_keywords,
            "unmatched": unmatched_keywords
        },
        "reasoning": {
            "Keyword Overlap": keyword_reasoning,
            "Formatting Issues": formatting_reasoning,
            "Missing Sections": missing_section_reasoning,
            "Conclusion": conclusion_reasoning,
        }
    }


# ---------- Resume Feedback Agent ----------
resume_feedback = Agent(
    role="Professional Resume Advisor",
    goal="Give feedback on the resume to make it stand out in the job market.",
    verbose=True,
    backstory="With a strategic mind and an eye for detail, you excel at providing feedback on resumes to highlight the most relevant skills and experiences."
)

resume_feedback_task = Task(
    description=(
        """Give feedback on the resume to make it stand out for recruiters. 
        Review every section, including the summary, work experience, skills, and education. 
        Suggest to add relevant sections if they are missing.  
        Also give an overall score to the resume out of 10.  
        This is the resume: {resume}"""
    ),
    expected_output="The overall score of the resume followed by the feedback in bullet points.",
    agent=resume_feedback
)

def get_resume_feedback(resume_text):
    from crewai import Crew, Task

    task = Task(
        description=(
            f"""You are a professional career coach and resume reviewer. 
Your task is to evaluate the resume critically and provide personalized, detailed feedback using the following format:
with proper formatting and nothing else

---

### 🛠️ **Areas to Improve**

1. **Formatting & Layout**  
Comment on section spacing, bullet alignment, and visual hierarchy.

2. **Content & Clarity**  
Point out where content is vague, repetitive, or could be improved with stronger verbs or clearer impact.

3. **Grammar / Typos / Terminology**  
Find specific errors (e.g., tool names, inconsistent abbreviations), and suggest clean alternatives.

---

### 📌 **Strengths**  
List 3–4 specific things done well — focus on concrete details from the resume, like quantified impact, strong tools/technologies, or clear structure.
---

### 💡 **Suggestions for Additional Sections**

- Suggest optional additions like a short summary/profile, links to GitHub or portfolio, or leadership/soft skill highlights.
- If anything is missing from the resume (e.g., GPA, project outcomes, relevant coursework), call that out.

Use markdown formatting (like bold and bullet points) to keep the output readable. Focus on *this specific resume* — not general advice.

Here is the resume to review:
{resume_text}
"""
        ),
        expected_output="Well-structured, specific feedback using the format above.",
        agent=resume_feedback
    )

    crew = Crew(
        agents=[resume_feedback],
        tasks=[task],
        verbose=True
    )

    return crew.kickoff()

from io import BytesIO

def extract_text_from_pdf(file_bytes):
    buffer = BytesIO(file_bytes)
    doc = fitz.open(stream=buffer.read(), filetype="pdf")
    return "".join([page.get_text() for page in doc])

def extract_text_from_docx(file_bytes):
    buffer = BytesIO(file_bytes)
    doc = docx.Document(buffer)
    return "\n".join(para.text for para in doc.paragraphs)

def extract_resume_text(file_bytes, file_name):
    if file_name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif file_name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    return "Unsupported file format."

#---------- Keyword Extraction ----------
def extract_keywords(text, limit=10):
    words = re.findall(r'\b\w+\b', text.lower())
    stopwords = {
        "and", "with", "the", "for", "to", "a", "in", "of", "on", "is", "by",
        "as", "an", "at", "from", "you", "your", "skills", "experience", "using",
        "performance", "projects", "github", "resume", "based", "etc", "work"
    }
    filtered = [word for word in words if len(word) > 2 and word not in stopwords]
    return [word for word, _ in Counter(filtered).most_common(limit)]



# ---------- OpenAI Embedding ----------
def embed_text(text):
    response = openai_client.embeddings.create(
        input=text,
        model="text-embedding-ada-002"
    )
    return response.data[0].embedding


def get_job_recommendations(resume_text):
    import random, re, requests, os
    from datetime import datetime, timezone
    from collections import Counter
    from sklearn.metrics.pairwise import cosine_similarity
    from docx import Document
    from docx.shared import Pt

    def generate_cover_letter_for_job(resume_text, job):
        prompt = f"""
        You're a professional resume writer. Generate a concise, tailored cover letter.

        Resume:
        {resume_text}

        Job Title: {job.get("job_title", "Unknown")}
        Company: {job.get("employer_name", "Unknown")}
        Job Description:
        {job.get("job_description", "")[:3000]}

        Guidelines:
        - Tailor it to the job.
        - Mention relevant experience/tools/skills.
        - Don't include placeholders like [Company Name]
        - Include today's date: {datetime.now().strftime('%B %d, %Y')}
        """
        try:
            response = openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=700
            )
            return response.choices[0].message.content.strip(), None
        except Exception as e:
            return None, f"Error generating cover letter: {e}"

    keywords = extract_keywords(resume_text, limit=10)
    query_keywords = [kw for kw in keywords if kw.isalpha()]

    if not query_keywords:
        words = re.findall(r'\b[a-zA-Z]{3,}\b', resume_text.lower())
        freq_words = [word for word, _ in Counter(words).most_common(10)]
        query_keywords = freq_words[:5]

    if not query_keywords:
        return [], []

    base_query = " ".join(query_keywords[:5])
    prefixes = ["", "now hiring", "top", "entry level", "urgent", "open roles"]
    suffixes = ["", "2025", "jobs", "remote", "analytics", "AI"]
    varied_query = f"{random.choice(prefixes)} {base_query} {random.choice(suffixes)}".strip()

    def fetch_jobs(query, pages=2):
        all_jobs = []
        for page in range(1, pages + 1):
            params = {
                "query": query,
                "location": "United States",
                "page": str(page),
                "num_pages": "1"
            }
            response = requests.get(
                "https://jsearch.p.rapidapi.com/search",
                headers={
                    "X-RapidAPI-Key": "66ab3e1826msh49f63b5561299a4p16d6c2jsn2f713ddce344",
                    "X-RapidAPI-Host": "jsearch.p.rapidapi.com"
                },
                params=params
            )
            if response.status_code == 200:
                all_jobs.extend(response.json().get("data", []))
        return all_jobs

    all_jobs = fetch_jobs(varied_query)
    if not all_jobs:
        all_jobs = fetch_jobs(base_query)
    if not all_jobs:
        return [], []

    seen = set()
    unique_jobs = []
    for job in all_jobs:
        key = (job.get("job_title"), job.get("employer_name"))
        if key not in seen:
            seen.add(key)
            unique_jobs.append(job)

    resume_words = set(re.findall(r'\b\w+\b', resume_text.lower()))

    def has_keyword_overlap(resume_words, job_desc, threshold=3):
        job_words = set(re.findall(r'\b\w+\b', job_desc.lower()))
        return len(resume_words & job_words) >= threshold

    filtered_jobs = [job for job in unique_jobs if has_keyword_overlap(resume_words, job.get("job_description", ""))]
    filtered_jobs = sorted(filtered_jobs, key=lambda j: len(j.get("job_description", "")), reverse=True)[:10]

    non_jd_df = plain_resumes_df[~plain_resumes_df["Resume File"].str.startswith("jd_")]
    if non_jd_df.empty:
        return [], []

    last_resume = non_jd_df.iloc[-1]
    resume_text = last_resume["Plain Text Resume"]

    try:
        resume_emb = embed_text(resume_text)
    except Exception:
        return [], []

    job_embedding_cache = {}
    scored = []
    for job in filtered_jobs:
        try:
            job_desc = job.get("job_description", "").strip()
            if not job_desc or len(job_desc) < 100:
                continue
            if job_desc in job_embedding_cache:
                job_emb = job_embedding_cache[job_desc]
            else:
                job_emb = embed_text(job_desc)
                job_embedding_cache[job_desc] = job_emb

            sim = cosine_similarity([resume_emb], [job_emb])[0][0]
            if sim > 0.5:
                scored.append((sim, job))
        except Exception:
            continue

    if not scored:
        return [], []

    top_jobs = sorted(scored, key=lambda x: x[0], reverse=True)[:5]

    def sanitize_filename(text):
        return re.sub(r'[^a-zA-Z0-9_\-]', '_', text)[:60]

    def save_cover_letter(text, filename):
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)
        for para in text.split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        os.makedirs("generated_letters", exist_ok=True)
        filepath = os.path.join("generated_letters", filename)
        doc.save(filepath)
        return filepath

    markdown_results = []
    cover_letter_paths = []

    for score, job in top_jobs:
        title = job.get("job_title", "N/A")
        company = job.get("employer_name", "N/A")
        location = job.get("job_city", "N/A")
        posted_datetime_str = job.get("job_posted_at_datetime_utc")

        if posted_datetime_str:
            try:
                posted_date = datetime.strptime(posted_datetime_str, "%Y-%m-%dT%H:%M:%S.%fZ").replace(
                    tzinfo=timezone.utc)
                days_ago = (datetime.now(timezone.utc) - posted_date).days
                posted = f"{days_ago} days ago" if days_ago > 0 else "Today"
            except:
                posted = "Recent"
        else:
            posted = "Recent"

        desc = job.get("job_description", "")[:200].replace("\n", " ").strip()
        desc = desc.rsplit(".", 1)[0] + "."

        match_percent = round(score * 100)
        apply_link = job.get("job_apply_link") or job.get("job_google_link") or "#"

        # Generate and save cover letter
        cover_letter_text, error = generate_cover_letter_for_job(resume_text, job)
        file_path = None

        if cover_letter_text and not error:
            filename = sanitize_filename(f"{title}_{company}.docx")
            file_path = save_cover_letter(cover_letter_text, filename)

        cover_letter_paths.append(file_path)

        markdown_results.append(f"""
### 🎯 [{title}]({apply_link}) - {company}
**Location**: {location}  
**Posted**: {posted}  
**Match Score**: {match_percent}%  
**Description**: {desc} 
""")

    # Clean markdown and filter valid paths
    clean_markdown = "\n---\n".join(
        job.replace('</div>', '').replace('</.div>', '').strip()
        for job in markdown_results
    )
    valid_paths = [path for path in cover_letter_paths if path is not None]

    return clean_markdown, valid_paths


# ---------- Chatbot Agent ----------
career_chatbot = Agent(
    role="Career Chat Assistant",
    goal="Help users with questions about resumes, jobs, applications, and career tips.",
    verbose=True,
    backstory="You're an experienced career assistant with deep knowledge of resumes, interview strategies, job search tips, and professional guidance. You respond in a friendly and concise manner."
)

def get_chatbot_response(user_query):
    chatbot_task = Task(
        description=f"""You are a career assistant chatbot helping job seekers.
Answer the following user question clearly, concisely, and accurately:

{user_query}
""",
        expected_output="Helpful, friendly response to user's question.",
        agent=career_chatbot
    )

    crew = Crew(
        agents=[career_chatbot],
        tasks=[chatbot_task],
        verbose=False
    )

    return crew.kickoff()

def generate_cover_letter():
    import os
    from docx import Document
    from docx.shared import Pt

    try:
        non_jd_df = plain_resumes_df[~plain_resumes_df["Resume File"].str.startswith("jd_")]
        resume_row = non_jd_df.iloc[-1]
        resume_text = resume_row["Plain Text Resume"]
        resume_file = resume_row["Metadata"]["resume_file"]
        jd_row = plain_resumes_df[plain_resumes_df["Resume File"] == f"jd_{resume_file}"]
        if jd_row.empty:
            return f"❌ No matching job description found for resume `{resume_file}`.", None
        jd_text = jd_row.iloc[0]["Plain Text Resume"]
    except Exception as e:
        return f"❌ Error retrieving resume and JD text: {e}", None

    # === Step 1: Extract personal details ===
    extract_prompt = f"""
You're an intelligent information extractor.

From the following resume text, extract:
1. **Full Name**
2. **Email Address**
3. **Phone Number** (if available)
4. **Location** (City, State or Country)

Return in the following format:
Name: ...
Email: ...
Phone: ...
Location: ...

Resume:
{resume_text}
    """
    try:
        extract_response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": extract_prompt}],
            temperature=0.2,
            max_tokens=300
        )
        extracted_info = extract_response.choices[0].message.content.strip()
        name = re.search(r"Name:\s*(.*)", extracted_info)
        email = re.search(r"Email:\s*(.*)", extracted_info)
        phone = re.search(r"Phone:\s*(.*)", extracted_info)
        location = re.search(r"Location:\s*(.*)", extracted_info)

        user_name = name.group(1).strip() if name else "Candidate"
        user_email = email.group(1).strip() if email else ""
        user_phone = phone.group(1).strip() if phone else ""
        user_location = location.group(1).strip() if location else ""
    except Exception as e:
        user_name, user_email, user_phone, user_location = "Candidate", "", "", ""

    # === Step 2: Analyze resume to get a summary ===
    summary_prompt = f"""
Here is the content of the resume:

{resume_text}

Please analyze the content and provide:

1. **Key Skills** (as a bullet list)
2. **Work Experience Summary**
3. **Recommended Job Title**
    """
    try:
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a resume expert and job matcher."},
                {"role": "user", "content": summary_prompt}
            ],
            temperature=0.3,
            max_tokens=700
        )
        resume_summary = response.choices[0].message.content.strip()
    except Exception as e:
        return f"❌ Error during resume analysis: {e}", None

    # === Step 3: Generate the cover letter ===
    cover_prompt = f"""
Using the following details, write a personalized and concise cover letter for the candidate.

{user_name}
{user_email}
{user_phone}
{user_location}

Resume Summary:
{resume_summary}

Job Description:
{jd_text}

Guidelines:
- Write as if it's a real cover letter
- Tailor it to the job
- No placeholders like [Company Name]
- Include contact details at the top
- Keep it professional, clear, and one page long
    """
    try:
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an expert in business writing."},
                {"role": "user", "content": cover_prompt}
            ],
            temperature=0.3,
            max_tokens=700
        )
        cover_letter_text = response.choices[0].message.content.strip()
    except Exception as e:
        return f"❌ Error generating cover letter: {e}", None

    # === Step 4: Save to .docx ===
    output_path = "sample_data/generated_cover_letter.docx"
    os.makedirs("sample_data", exist_ok=True)

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    for para in cover_letter_text.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.save(output_path)
    return cover_letter_text, output_path
