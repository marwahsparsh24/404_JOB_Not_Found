import os
import fitz
import chromadb
import pandas as pd
from pymongo import MongoClient
from sklearn.metrics.pairwise import cosine_similarity
import docx
import re
import requests
import numpy as np
from collections import Counter
from openai import OpenAI
from crewai import Agent, Task, Crew
from datetime import datetime, timezone
import os
os.environ["OPENAI_API_KEY"] = "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"
# 🔐 Load API key from environment variable
openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Define folder paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RESUME_DIR = os.path.join(BASE_DIR, 'data', 'resumes')
JD_DIR = os.path.join(BASE_DIR,'data','job_descriptions')
STATUS_DIR = os.path.join(BASE_DIR,'data','application_status')


def extract_formatted_text(file_path):
    """
    Extract formatted text from PDF or DOCX files.
    For PDFs: preserve font styles, sizes, layout using PyMuPDF.
    For DOCX: extract paragraph-wise text without formatting details.
    """
    try:
        if file_path.lower().endswith(".pdf"):
            doc = fitz.open(file_path)
            formatted_text = ""

            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                blocks = page.get_text("dict")["blocks"]

                for block in blocks:
                    if block["type"] == 0:  # Only process text blocks
                        for line in block["lines"]:
                            for span in line["spans"]:
                                text = span["text"]
                                font_size = span["size"]
                                font_color = span["color"]
                                font_name = span["font"]
                                is_bold = "bold" in font_name.lower()
                                is_italic = "italic" in font_name.lower()

                                formatted_text += (
                                    f"Text: {text}, Font: {font_name}, Size: {font_size}, "
                                    f"Color: {font_color}, Bold: {is_bold}, Italic: {is_italic}\n"
                                )
            return formatted_text

        elif file_path.lower().endswith(".docx"):
            doc = docx.Document(file_path)
            formatted_text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    formatted_text += f"Text: {para.text}, Font: N/A, Size: N/A, Color: N/A, Bold: N/A, Italic: N/A\n"
            return formatted_text

        else:
            return "Unsupported file format."

    except Exception as e:
        return f"Error extracting formatted text from {file_path}: {e}"


def extract_plain_text(pdf_path):
    """Extract plain text (ATS-friendly) using PyMuPDF."""
    try:
        doc = fitz.open(pdf_path)
        plain_text = ""

        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            plain_text += page.get_text("text")

        return plain_text
    except Exception as e:
        print(f"Error extracting plain text from {pdf_path}: {e}")
        return ""


def read_text_file(file_path):
    """Read text from JD or application status files."""
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read().strip()
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return ""


def process_resumes():
    """Processes resumes, JDs, and application statuses into a Pandas DataFrame."""
    data = []

    # List all files in the directories
    resume_files = sorted([file for file in os.listdir(RESUME_DIR) if file.lower().endswith('.pdf')])
    jd_files = sorted([file for file in os.listdir(JD_DIR) if file.lower().endswith('.txt')])
    status_files = sorted([file for file in os.listdir(STATUS_DIR) if file.lower().endswith('.txt')])

    # Ensure that the lengths of resume_files, jd_files, and status_files are consistent
    num_files = min(len(resume_files), len(jd_files), len(status_files))

    for i in range(num_files):
        resume_file = resume_files[i]
        jd_file = jd_files[i]
        status_file = status_files[i]

        resume_path = os.path.join(RESUME_DIR, resume_file)
        jd_path = os.path.join(JD_DIR, jd_file)
        status_path = os.path.join(STATUS_DIR, status_file)

        formatted_text = extract_formatted_text(resume_path)
        plain_text = extract_plain_text(resume_path)
        job_description = read_text_file(jd_path)
        application_status = read_text_file(status_path)

        # Store everything in a structured list
        data.append({
            "Resume File": resume_file,
            "Plain Text Resume": plain_text,
            "Formatted Resume": formatted_text,
            "Job Description": job_description,
            "Application Status": application_status
        })

    # Convert to DataFrame
    df = pd.DataFrame(data)

    return df

# Run the script and get DataFrame
df = process_resumes()
# Display DataFrame


mongo_client = MongoClient("mongodb://localhost:27017/")
db = mongo_client["resume_db"]
formatted_resume_collection = db["formatted_resumes"]

client = chromadb.PersistentClient(path="./chroma_db")
collection = client.get_or_create_collection("plain_resumes")


def store_to_chromadb_and_mongodb(df):
    for index, row in df.iterrows():
        resume_file = row["Resume File"]
        plain_text = row["Plain Text Resume"]
        formatted_text = row["Formatted Resume"]
        job_description = row["Job Description"]
        application_status = row["Application Status"]

        # MongoDB
        existing_resume = formatted_resume_collection.find_one({"resume_file": resume_file})
        if existing_resume:
            formatted_resume_collection.update_one(
                {"resume_file": resume_file},
                {"$set": {
                    "formatted_resume": formatted_text,
                    "job_description": job_description,
                    "application_status": application_status
                }}
            )
            print(f"Updated existing entry in MongoDB for: {resume_file}")
        else:
            formatted_resume_collection.insert_one({
                "resume_file": resume_file,
                "formatted_resume": formatted_text,
                "job_description": job_description,
                "application_status": application_status
            })
            print(f"Inserted new entry in MongoDB for: {resume_file}")

        # === Store Resume Embedding in ChromaDB ===
        resume_doc_id = f"resume_{resume_file}"
        if collection.get(ids=[resume_doc_id])['documents']:
            collection.delete(ids=[resume_doc_id])
        collection.add(
            documents=[plain_text],
            metadatas=[{
                "type": "resume",
                "resume_file": resume_file,
                "job_description": job_description,
                "application_status": application_status
            }],
            ids=[resume_doc_id]
        )
        print(f"Stored in ChromaDB: {resume_doc_id}")

        # === Store JD Embedding in ChromaDB ===
        jd_doc_id = f"jd_{resume_file}"
        if collection.get(ids=[jd_doc_id])['documents']:
            collection.delete(ids=[jd_doc_id])
        collection.add(
            documents=[job_description],
            metadatas=[{
                "type": "job_description",
                "resume_file": resume_file,
                "application_status": application_status
            }],
            ids=[jd_doc_id]
        )
        print(f"Stored in ChromaDB: {jd_doc_id}")

store_to_chromadb_and_mongodb(df)

# MongoDB connection string (adjust it as needed)
MONGO_URI = "mongodb://localhost:27017"
DATABASE_NAME = "resume_db"
FORMATTED_COLLECTION = "formatted_resumes"

# Connect to MongoDB
client = MongoClient(MONGO_URI)
db = client[DATABASE_NAME]
formatted_resume_collection = db[FORMATTED_COLLECTION]


def retrieve_formatted_resumes_as_df():
    """Retrieve unique formatted resume data from MongoDB and return it as a DataFrame"""

    # Retrieve all documents from the collection
    formatted_resumes = formatted_resume_collection.find()

    # Convert the results into a list of dictionaries while ensuring unique entries
    resume_list = []
    seen_files = set()  # To track unique resume files

    for resume in formatted_resumes:
        resume_file = resume['resume_file']

        # Only append if the resume_file hasn't been encountered before
        if resume_file not in seen_files:
            seen_files.add(resume_file)
            resume_list.append({
                "Resume File": resume['resume_file'],
                "Formatted Resume": resume['formatted_resume'],
                "Job Description": resume['job_description'],
                "Application Status": resume['application_status']
            })

    # Convert the list of dictionaries to a DataFrame
    df = pd.DataFrame(resume_list)
    return df


# Example usage: retrieve unique formatted resumes as a DataFrame
formatted_resume_df = retrieve_formatted_resumes_as_df()
print(formatted_resume_df)
# Initialize ChromaDB client and collection
client = chromadb.PersistentClient(path="./chroma_db")  # Use the correct path if needed
collection = client.get_or_create_collection(name="plain_resumes")


def retrieve_all_resumes():
    """Retrieve all resumes stored in ChromaDB and return as a DataFrame"""
    results = collection.get()  # Retrieves all stored documents

    # Convert results to a structured DataFrame
    resume_list = []
    for i in range(len(results['documents'])):
        resume_list.append({
            "Resume File": results['ids'][i],
            "Plain Text Resume": results['documents'][i],
            "Metadata": results['metadatas'][i]  # Contains metadata like resume_file, job_description, etc.
        })

    df = pd.DataFrame(resume_list)
    return df


# Retrieve all resumes
plain_resumes_df = retrieve_all_resumes()


INPUT_DIR = os.path.join(BASE_DIR,'data', 'inputdata')
os.makedirs(INPUT_DIR, exist_ok=True)

# Initialize empty DataFrames if not already defined
try:
    formatted_resume_df
except NameError:
    formatted_resume_df = pd.DataFrame(columns=["Resume File", "Formatted Resume", "Job Description", "Application Status"])

try:
    plain_resumes_df
except NameError:
    plain_resumes_df = pd.DataFrame(columns=["Resume File", "Plain Text Resume", "Metadata"])

# Save uploaded file to disk (used by backend functions)

def save_uploaded_file(file_name, file_bytes):
    """Save uploaded resume to disk in the input folder."""
    file_path = os.path.join(INPUT_DIR, file_name)
    with open(file_path, "wb") as f:
        f.write(file_bytes)
    return file_name  # Just returning the filename

# === Core Function ===
def upload_rejected_resume_local_to_df(resume_filename, jd_text, application_status):
    if application_status.lower() != "rejected":
        return {"error": "Application status must be 'rejected'"}

    if not jd_text.strip():
        return {"error": "Job description text cannot be empty."}

    resume_path = os.path.join(INPUT_DIR, resume_filename)

    formatted_text = extract_formatted_text(resume_path)
    plain_text = extract_plain_text(resume_path)

    formatted_row = {
        "Resume File": resume_filename,
        "Formatted Resume": formatted_text,
        "Job Description": jd_text,
        "Application Status": application_status
    }

    plain_row = {
        "Resume File": f"resume_{resume_filename}",
        "Plain Text Resume": plain_text,
        "Metadata": {
            "application_status": application_status,
            "job_description": jd_text,
            "resume_file": resume_filename,
        }
    }

    jd_row = {
        "Resume File": f"jd_{resume_filename}",
        "Plain Text Resume": jd_text,
        "Metadata": {
            "application_status": application_status,
            "resume_file": resume_filename,
            "type": "job_description"
        }
    }

    global formatted_resume_df, plain_resumes_df
    formatted_resume_df = formatted_resume_df[formatted_resume_df["Resume File"] != resume_filename]
    plain_resumes_df = plain_resumes_df[
        ~plain_resumes_df["Resume File"].isin([f"resume_{resume_filename}", f"jd_{resume_filename}"])]

    formatted_resume_df = pd.concat([formatted_resume_df, pd.DataFrame([formatted_row])], ignore_index=True)
    plain_resumes_df = pd.concat([plain_resumes_df, pd.DataFrame([plain_row, jd_row])], ignore_index=True)

    # === Store in MongoDB ===
    existing_resume = formatted_resume_collection.find_one({"resume_file": resume_filename})
    if existing_resume:
        formatted_resume_collection.update_one(
            {"resume_file": resume_filename},
            {"$set": {
                "formatted_resume": formatted_text,
                "job_description": jd_text,
                "application_status": application_status
            }}
        )
    else:
        formatted_resume_collection.insert_one({
            "resume_file": resume_filename,
            "formatted_resume": formatted_text,
            "job_description": jd_text,
            "application_status": application_status
        })

    # === Store in ChromaDB ===
    resume_doc_id = f"resume_{resume_filename}"
    jd_doc_id = f"jd_{resume_filename}"

    if collection.get(ids=[resume_doc_id])["documents"]:
        collection.delete(ids=[resume_doc_id])
    if collection.get(ids=[jd_doc_id])["documents"]:
        collection.delete(ids=[jd_doc_id])

    collection.add(
        documents=[plain_text],
        metadatas=[{
            "type": "resume",
            "resume_file": resume_filename,
            "job_description": jd_text,
            "application_status": application_status
        }],
        ids=[resume_doc_id]
    )

    collection.add(
        documents=[jd_text],
        metadatas=[{
            "type": "job_description",
            "resume_file": resume_filename,
            "application_status": application_status
        }],
        ids=[jd_doc_id]
    )

    return {
        "message": "Files processed and stored in local DataFrames, MongoDB, and ChromaDB (replaced if existed)",
        "resume_file": resume_filename
    }


# === 1. FORMATTING COMPLIANCE CHECK ===
def run_formatting_compliance_last():
    global formatted_resume_df

    if formatted_resume_df.empty:
        return "⚠️ DataFrame is empty."

    last_row = formatted_resume_df.iloc[-1]
    formatted_text = last_row["Formatted Resume"]
    resume_file = last_row["Resume File"]

    system_prompt = (
        "You are an expert in resume formatting for Applicant Tracking Systems (ATS). "
        "Your task is to analyze formatting features extracted from resumes and identify ATS-related issues."
    )

    user_prompt = f'''
Here is the formatting output extracted from a resume:

{formatted_text}

You are an ATS compliance reviewer. Check the resume against the following formatting rules and flag any violations:

1. Use a single-column layout only (no multi-columns or tables).
2. Use readable fonts.
3. Font size should be between 10pt and 14pt.
4. Avoid decorative or unusual fonts, symbols or special characters.
5. Do not use white/invisible text.
6. Do not include icons, graphics, logos, headshots, or image-based elements.
7. Avoid using tables (prefer plain text and bullet points instead).
8. Resume should be a text-based PDF or DOCX file, not an image-based scan.

✅ If no issues are found, respond with:  
'✅ No major ATS formatting issues found.'

⚠️ Do NOT flag the `•` bullet as a formatting issue.
'''

    try:
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.2,
            max_tokens=500
        )
        report = response.choices[0].message.content.strip()
    except Exception as e:
        report = f"Error during analysis: {e}"

    print(f"✅ Formatting check completed for: {resume_file}")
    return report

# === 2. ATS PARSING ERROR DETECTION ===
def ats_parsing_error_detection_agent():
    global plain_resumes_df

    if plain_resumes_df.empty:
        return "⚠️ DataFrame is empty."

    non_jd_df = plain_resumes_df[~plain_resumes_df["Resume File"].str.startswith("jd_")]
    if non_jd_df.empty:
        return "⚠️ No resume entries found (JD rows excluded)."

    last_row = non_jd_df.iloc[-1]
    plain_text = last_row["Plain Text Resume"]
    resume_file = last_row["Resume File"]

    system_prompt = (
        "You are an ATS parsing validator. Your job is to simulate how an ATS system would interpret a resume "
        "and identify issues that may cause parsing errors or misinterpretation."
    )

    user_prompt = f"""
Here is the plain text of a resume:

--------------------
{plain_text}
--------------------

Analyze the resume and check for the following:

1. **Section Header Validation**
   - Are standard sections like Contact Info, Summary, Skills, Experience, Education present?
   - Are any headers mislabeled or non-standard?

2. **Content Completeness**
   - Are sections populated with useful content (not just headers)?
   - Flag incomplete sections.

3. **Text Extraction Quality**
   - Are there signs of parsing errors (garbled formatting, symbols, broken lines, missing content)?

Return a bullet-point summary of issues found. If no issues, say:  
✅ No major parsing or content issues found.
"""

    try:
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3,
            max_tokens=600
        )
        report = response.choices[0].message.content.strip()
    except Exception as e:
        report = f"Error during analysis: {e}"

    print(f"✅ ATS parsing check complete for: {resume_file}")
    return report

# ## ATS Scoring Agent
# Function to generate OpenAI embedding
def embed_text(text):
    response = openai_client.embeddings.create(
        input=text,
        model="text-embedding-ada-002"
    )
    return response.data[0].embedding


# ATS Scoring Agent
def run_ats_scoring_with_embedding():
    # Filter only resumes
    non_jd_df = plain_resumes_df[~plain_resumes_df["Resume File"].str.startswith("jd_")]
    if non_jd_df.empty:
        return "⚠️ No resumes to score."

    last_resume = non_jd_df.iloc[-1]
    resume_file = last_resume["Resume File"]
    resume_text = last_resume["Plain Text Resume"]
    jd_entry = plain_resumes_df[plain_resumes_df["Resume File"] == f"jd_{resume_file.split('_', 1)[-1]}"]

    if jd_entry.empty:
        return f"❌ No matching job description found for resume `{resume_file}`."

    jd_text = jd_entry.iloc[0]["Plain Text Resume"]

    # Generate embeddings using OpenAI
    resume_embedding = embed_text(resume_text)
    jd_embedding = embed_text(jd_text)

    if resume_embedding is None or jd_embedding is None:
        return "❌ Failed to generate embeddings."

    # Cosine similarity
    sim = cosine_similarity([resume_embedding], [jd_embedding])[0][0]
    ats_score = round(sim * 100, 2)

    # LLM Reasoning
    try:
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an expert in evaluating resumes against job descriptions."},
                {"role": "user", "content": f"""
Given the following resume and job description:

=== Resume ===
{resume_text}

=== Job Description ===
{jd_text}

Please provide a short explanation on how well this resume matches the job description in terms of relevant skills, experience, and keywords. Highlight any major mismatches or missing areas.
"""}
            ],
            temperature=0.3,
            max_tokens=500
        )
        reasoning = response.choices[0].message.content.strip()
    except Exception as e:
        reasoning = f"(LLM reasoning failed: {e})"

    return f"""
📊 **ATS Score** for `{resume_file}`: {ats_score}%

🧠 **LLM Reasoning:**
{reasoning}
"""


# ---------- Resume Feedback Agent ----------
resume_feedback = Agent(
    role="Professional Resume Advisor",
    goal="Give feedback on the resume to make it stand out in the job market.",
    verbose=True,
    backstory="With a strategic mind and an eye for detail, you excel at providing feedback on resumes to highlight the most relevant skills and experiences."
)

resume_feedback_task = Task(
    description=(
        """Give feedback on the resume to make it stand out for recruiters. 
        Review every section, including the summary, work experience, skills, and education. 
        Suggest to add relevant sections if they are missing.  
        Also give an overall score to the resume out of 10.  
        This is the resume: {resume}"""
    ),
    expected_output="The overall score of the resume followed by the feedback in bullet points.",
    agent=resume_feedback
)

def get_resume_feedback(resume_text):
    from crewai import Crew, Task

    task = Task(
        description=(
            f"""You are a professional career coach and resume reviewer. 
Your task is to evaluate the resume critically and provide personalized, detailed feedback using the following format:

---

## ✅ **Overall Resume Score**: X/10  
Give a reason for the score, based on technical depth, relevance, formatting, and clarity.

---

### 📌 **Strengths**  
List 3–4 specific things done well — focus on concrete details from the resume, like quantified impact, strong tools/technologies, or clear structure.

---

### 🛠️ **Areas to Improve**

1. **Formatting & Layout**  
Comment on section spacing, bullet alignment, and visual hierarchy.

2. **Content & Clarity**  
Point out where content is vague, repetitive, or could be improved with stronger verbs or clearer impact.

3. **Grammar / Typos / Terminology**  
Find specific errors (e.g., tool names, inconsistent abbreviations), and suggest clean alternatives.

---

### 💡 **Suggestions for Additional Sections**

- Suggest optional additions like a short summary/profile, links to GitHub or portfolio, or leadership/soft skill highlights.
- If anything is missing from the resume (e.g., GPA, project outcomes, relevant coursework), call that out.

Use markdown formatting (like bold and bullet points) to keep the output readable. Focus on *this specific resume* — not general advice.

Here is the resume to review:
{resume_text}
"""
        ),
        expected_output="Well-structured, specific feedback using the format above.",
        agent=resume_feedback
    )

    crew = Crew(
        agents=[resume_feedback],
        tasks=[task],
        verbose=True
    )

    return crew.kickoff()
#
# # ---------- Resume Text Extraction ----------
# def extract_text_from_pdf(file):
#     doc = fitz.open(stream=file.read(), filetype="pdf")
#     return "".join([page.get_text() for page in doc])
#
# def extract_text_from_docx(file):
#     doc = docx.Document(file)
#     return "\n".join(para.text for para in doc.paragraphs)
#
# def extract_resume_text(uploaded_file):
#     if uploaded_file.name.endswith(".pdf"):
#         return extract_text_from_pdf(uploaded_file)
#     elif uploaded_file.name.endswith(".docx"):
#         return extract_text_from_docx(uploaded_file)
#     return "Unsupported file format."

from io import BytesIO

def extract_text_from_pdf(file_bytes):
    buffer = BytesIO(file_bytes)
    doc = fitz.open(stream=buffer.read(), filetype="pdf")
    return "".join([page.get_text() for page in doc])

def extract_text_from_docx(file_bytes):
    buffer = BytesIO(file_bytes)
    doc = docx.Document(buffer)
    return "\n".join(para.text for para in doc.paragraphs)

def extract_resume_text(file_bytes, file_name):
    if file_name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif file_name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    return "Unsupported file format."

#---------- Keyword Extraction ----------
def extract_keywords(text, limit=10):
    words = re.findall(r'\b\w+\b', text.lower())
    stopwords = {
        "and", "with", "the", "for", "to", "a", "in", "of", "on", "is", "by",
        "as", "an", "at", "from", "you", "your", "skills", "experience", "using",
        "performance", "projects", "github", "resume", "based", "etc", "work"
    }
    filtered = [word for word in words if len(word) > 2 and word not in stopwords]
    return [word for word, _ in Counter(filtered).most_common(limit)]



# ---------- OpenAI Embedding ----------
def embed_text(text):
    response = openai_client.embeddings.create(
        input=text,
        model="text-embedding-ada-002"
    )
    return response.data[0].embedding



# ---------- Job Recommendations ----------
def get_job_recommendations(resume_text):
    import random

    keywords = extract_keywords(resume_text, limit=10)
    # ✅ Filter: only keep alphabetic keywords (no UUIDs, numbers, or symbols)
    query_keywords = [kw for kw in keywords if kw.isalpha()]

    # ✅ If nothing usable, try with the most frequent longer words
    if not query_keywords:
        words = re.findall(r'\b[a-zA-Z]{3,}\b', resume_text.lower())
        freq_words = [word for word, _ in Counter(words).most_common(10)]
        query_keywords = freq_words[:5]

    # Still if nothing found, return a clean error
    if not query_keywords:
        return "❌ Could not extract meaningful keywords from resume."

    base_query = " ".join(query_keywords[:5])
    prefixes = ["", "now hiring", "top", "entry level"]
    suffixes = ["", "2024", "jobs", "remote"]
    varied_query = f"{random.choice(prefixes)} {base_query} {random.choice(suffixes)}".strip()

    print(f"🧠 Querying with: {varied_query}")

    def fetch_jobs(query):
        all_jobs = []
        print(f"🔍 Sending query to JSearch: '{query}'")
        for page in range(1, 3):
            params = {
                "query": query,
                "location": "United States",
                "page": str(page),
                "num_pages": "1"
            }
            response = requests.get(
                "https://jsearch.p.rapidapi.com/search",
                headers={
                    "X-RapidAPI-Key": "17120e94c0msh55896932e9683a0p1252cejsna9e96373e0b2",
                    "X-RapidAPI-Host": "jsearch.p.rapidapi.com"
                },
                params=params
            )
            if response.status_code == 200:
                jobs = response.json().get("data", [])
                print(f"✅ Page {page} returned {len(jobs)} jobs")
                all_jobs.extend(jobs)
            else:
                print(f"❌ API error: {response.status_code} - {response.text}")
        return all_jobs

    all_jobs = fetch_jobs(varied_query)
    if not all_jobs:
        print("⚠️ No jobs found with varied query. Trying base query...")
        all_jobs = fetch_jobs(base_query)

    if not all_jobs:
        return "❌ No jobs found. Try simplifying your resume or keywords."

    seen = set()
    unique_jobs = []
    for job in all_jobs:
        key = (job.get("job_title"), job.get("employer_name"))
        if key not in seen:
            seen.add(key)
            unique_jobs.append(job)

    random.shuffle(unique_jobs)
    sampled_jobs = random.sample(unique_jobs, min(25, len(unique_jobs)))

    try:
        resume_emb = embed_text(resume_text)
    except Exception as e:
        return f"❌ Embedding error: {e}"

    scored = []
    for job in sampled_jobs:
        try:
            job_desc = job.get("job_description", "").strip()
            if not job_desc:
                continue
            job_emb = embed_text(job_desc)
            # Filter only resumes
            non_jd_df = plain_resumes_df[~plain_resumes_df["Resume File"].str.startswith("jd_")]
            if non_jd_df.empty:
                return "⚠️ No resumes to score."

            last_resume = non_jd_df.iloc[-1]
            resume_file = last_resume["Resume File"]
            resume_text = last_resume["Plain Text Resume"]
            jd_entry = plain_resumes_df[plain_resumes_df["Resume File"] == f"jd_{resume_file.split('_', 1)[-1]}"]

            if jd_entry.empty:
                return f"❌ No matching job description found for resume `{resume_file}`."

            jd_text = jd_entry.iloc[0]["Plain Text Resume"]

            # Generate embeddings using OpenAI
            resume_embedding = embed_text(resume_text)
            jd_embedding = embed_text(jd_text)

            if resume_embedding is None or jd_embedding is None:
                return "❌ Failed to generate embeddings."
            sim = cosine_similarity([resume_embedding], [jd_embedding])[0][0]
            ats_score = round(sim * 100, 2)
            scored.append((sim, job))
        except Exception:
            continue

    top_jobs = sorted(scored, key=lambda x: x[0], reverse=True)[:5]
    if not top_jobs:
        return "❌ No relevant jobs found."

    from datetime import datetime, timezone

    results = []
    for score, job in top_jobs:
        title = job.get("job_title", "N/A")
        company = job.get("employer_name", "N/A")
        location = job.get("job_city", "N/A")
        posted_datetime_str = job.get("job_posted_at_datetime_utc")

        if posted_datetime_str:
            posted_date = datetime.strptime(posted_datetime_str, "%Y-%m-%dT%H:%M:%S.%fZ").replace(tzinfo=timezone.utc)
            days_ago = (datetime.now(timezone.utc) - posted_date).days
            posted = f"{days_ago} days ago" if days_ago > 0 else "Today"
        else:
            posted = "Recent"

        desc = job.get("job_description", "")[:200].replace("\n", " ").strip()
        desc = desc.rsplit(".", 1)[0] + "."
        match_percent = round(score * 100)
        apply_link = job.get("job_apply_link") or job.get("job_google_link", "#")

        results.append(f"""
    **[{title}]({apply_link}) - {company}**  
    Location: {location}  
    Posted: {posted}  
    Match Score: {match_percent}%  
    Description: {desc}
    """)

    return "\n\n".join(results)


